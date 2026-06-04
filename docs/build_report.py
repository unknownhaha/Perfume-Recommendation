"""
Generates docs/report.docx — technical report for the Perfume Recommendation project.
Run: python docs/build_report.py

Changes from v1:
 - All file/variable names replaced with conceptual descriptions (academic style)
 - Section 2.4 uses a real matplotlib block diagram
 - Section 1.3.2 has properly formatted Cosine + Jaccard equations
 - Section 3.2 explains Avg Score formula
 - In-text citations added for [8] Salton and [9] Jaccard
 - Appendix retains technical file names (appropriate there)
"""
import io
import os
import shutil

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(DOCS_DIR, "report.docx")
DIAG_PATH = os.path.join(DOCS_DIR, "_pipeline_diagram.png")
EQ_DIR = os.path.join(DOCS_DIR, "_eq_cache")
os.makedirs(EQ_DIR, exist_ok=True)

BLACK = (0, 0, 0)
FONT_BODY = "TH Sarabun New"
FONT_SIZE_BODY = 16
FONT_SIZE_H1 = 20
FONT_SIZE_H2 = 18
FONT_SIZE_H3 = 16

# ── diagram generator ────────────────────────────────────────────────────────

def make_pipeline_diagram(out_path):
    fig, axes = plt.subplots(1, 2, figsize=(14, 7))
    fig.patch.set_facecolor("white")

    BOX_FILL = "#FFFFFF"
    BOX_EDGE = "#000000"
    ARROW = "#000000"

    def draw_box(ax, x, y, w, h, label, sublabel="", color=BOX_FILL, text_color="#000000",
                 bold=False, radius=0.06):
        box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                             boxstyle=f"round,pad=0.02,rounding_size={radius}",
                             facecolor=color, edgecolor=BOX_EDGE, linewidth=1.2)
        ax.add_patch(box)
        fs = 9.5 if bold else 9
        fw = "bold" if bold else "normal"
        ax.text(x, y + (0.07 if sublabel else 0), label,
                ha="center", va="center", fontsize=fs, fontweight=fw,
                color=text_color, wrap=True)
        if sublabel:
            ax.text(x, y - 0.17, sublabel,
                    ha="center", va="center", fontsize=7.5,
                    color="#000000", style="italic")

    def arrow(ax, x1, y1, x2, y2):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=ARROW, lw=1.4))

    # ── Left: Offline Pipeline ─────────────────────────────────────────
    ax = axes[0]
    ax.set_xlim(0, 1)
    ax.set_ylim(-0.3, 6.8)
    ax.axis("off")
    ax.set_facecolor("white")
    ax.set_title("Offline Pipeline (Build Stage)", fontsize=12,
                 fontweight="bold", color="#000000", pad=10)

    steps_off = [
        (0.5, 6.2, "Raw Dataset", "~26,000 perfumes (JSON)", BOX_FILL, "#000000", True),
        (0.5, 5.0, "Data Cleaning & Normalization",
         "Uppercase categories, Title-case ingredients", BOX_FILL, "#000000", False),
        (0.5, 3.8, "Vocabulary Filtering",
         "Remove rare notes (< 5 occurrences)\n394 → 312 unique notes", BOX_FILL, "#000000", False),
        (0.5, 2.6, "Feature Encoding",
         "Multi-Hot (ingredients) + One-Hot (family, subfamily, gender)", BOX_FILL, "#000000", False),
        (0.5, 1.4, "Feature Weighting",
         "Apply per-group weights (tuned via sweep)", BOX_FILL, "#000000", False),
        (0.5, 0.2, "Model Artifact Storage",
         "Encoders + Weighted Sparse Matrix + Config", BOX_FILL, "#000000", True),
    ]
    for (x, y, lbl, sub, c, tc, bd) in steps_off:
        draw_box(ax, x, y, 0.78, 0.6, lbl, sub, color=c, text_color=tc, bold=bd)
    for i in range(len(steps_off) - 1):
        _, y1, *_ = steps_off[i]
        _, y2, *_ = steps_off[i + 1]
        arrow(ax, 0.5, y1 - 0.3, 0.5, y2 + 0.3)

    # ── Right: Online Pipeline ─────────────────────────────────────────
    ax = axes[1]
    ax.set_xlim(0, 1)
    ax.set_ylim(-0.3, 6.8)
    ax.axis("off")
    ax.set_facecolor("white")
    ax.set_title("Online Pipeline (Inference Stage)", fontsize=12,
                 fontweight="bold", color="#000000", pad=10)

    steps_on = [
        (0.5, 6.2, "User Input",
         "Preferred notes, family, gender, brand", BOX_FILL, "#000000", True),
        (0.5, 5.0, "Query Vector Encoding",
         "Same encoding + weights as training", BOX_FILL, "#000000", False),
        (0.5, 3.8, "Catalog Filtering",
         "Filter by gender and brand", BOX_FILL, "#000000", False),
        (0.5, 2.6, "Cosine Similarity",
         "One sparse matrix pass over filtered catalog", BOX_FILL, "#000000", False),
        (0.5, 1.4, "Hybrid Re-ranking",
         "hybrid = 0.7 × cosine + 0.3 × Jaccard\n(on top-N×5 shortlist)", BOX_FILL, "#000000", False),
        (0.5, 0.2, "Top-N Recommendations",
         "Ranked results + Matched Notes explanation", BOX_FILL, "#000000", True),
    ]
    for (x, y, lbl, sub, c, tc, bd) in steps_on:
        draw_box(ax, x, y, 0.78, 0.6, lbl, sub, color=c, text_color=tc, bold=bd)
    for i in range(len(steps_on) - 1):
        _, y1, *_ = steps_on[i]
        _, y2, *_ = steps_on[i + 1]
        arrow(ax, 0.5, y1 - 0.3, 0.5, y2 + 0.3)

    plt.tight_layout(rect=[0, 0.03, 1, 0.97])
    plt.savefig(out_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()


make_pipeline_diagram(DIAG_PATH)

# ── doc helpers ──────────────────────────────────────────────────────────────

def set_font(run, name=FONT_BODY, size=FONT_SIZE_BODY, bold=False,
             color=BLACK, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rPr.insert(0, rFonts)


def add_heading(doc, text, size=FONT_SIZE_H1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    return p


def add_subheading(doc, text, level=2):
    sizes = {2: FONT_SIZE_H2, 3: FONT_SIZE_H3}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, FONT_SIZE_H3), bold=True)
    return p


def add_body(doc, text, first_line_indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(24)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    run = p.add_run(text)
    set_font(run, size=FONT_SIZE_BODY)
    return p


def render_equation_image(latex, filename, fontsize=16):
    """Render LaTeX math as a high-quality PNG using matplotlib mathtext."""
    path = os.path.join(EQ_DIR, filename)
    fig = plt.figure(figsize=(7.5, 0.9))
    fig.patch.set_facecolor("white")
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis("off")
    ax.text(0.5, 0.5, f"${latex}$", ha="center", va="center", fontsize=fontsize, color="black")
    plt.savefig(path, dpi=250, bbox_inches="tight", pad_inches=0.08, facecolor="white")
    plt.close()
    return path


def add_equation(doc, label, latex, width=Inches(5.5)):
    """Insert a rendered LaTeX equation image (Word-native math appearance)."""
    img_path = render_equation_image(latex, f"eq_{label}.png")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=width)
    if label:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(0)
        caption.paragraph_format.space_after = Pt(8)
        set_font(caption.add_run(f"({label})"), size=14, italic=True)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_font(run, size=15)
    return p


def add_numbered(doc, text, num):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{num}. {text}")
    set_font(run, size=15)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, name="Courier New", size=12)
    p.paragraph_format.line_spacing = Pt(16)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_font(run, size=14, italic=True)
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=14, bold=True)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "FFFFFF")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=14, bold=(ci == 0))
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "FFFFFF")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def page_break(doc):
    doc.add_page_break()


# ── build document ───────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# ── Cover page (KMUTT-style layout) ──────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
set_font(p.add_run("ระบบแนะนำน้ำหอมด้วยการกรองตามเนื้อหา"), size=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Content-Based Perfume Recommendation System"), size=18, bold=True)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("_______________________________"), size=16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("โครงงานนี้เป็นส่วนหนึ่งของรายวิชา Machine Learning  ปีการศึกษา 2568"), size=16)

doc.add_paragraph()

add_heading(doc, "บทคัดย่อ", size=FONT_SIZE_H1)
add_body(doc,
    "โครงงานนี้นำเสนอระบบแนะนำน้ำหอมแบบ Content-Based Filtering "
    "สำหรับชุดข้อมูลน้ำหอมประมาณ 26,000 รายการ จากแหล่งข้อมูล doevent/perfume บน HuggingFace "
    "ระบบสร้าง Feature Vector แบบ Sparse จากส่วนผสม (Ingredients) กลุ่มกลิ่น (Scent Family) "
    "และเพศ (Gender) โดยใช้การถ่วงน้ำหนักต่อกลุ่มคุณลักษณะที่ได้รับการจูนผ่านกระบวนการ Weight Sweep "
    "และจัดอันดับผลลัพธ์ด้วย Hybrid Score ที่ผสาน Cosine Similarity กับ Jaccard Overlap "
    "ผลการประเมินพบว่า Family Precision@10 เพิ่มขึ้น 16% และ Leave-One-Out Hit-Rate เพิ่มขึ้น 8% "
    "เมื่อเทียบกับระบบ Baseline ที่ใช้ Cosine Similarity เพียงอย่างเดียว "
    "ระบบไม่ต้องการข้อมูลคะแนนจากผู้ใช้ ตอบสนองได้เร็วน้อยกว่า 10 มิลลิวินาทีต่อการค้นหา "
    "และนำเสนอผ่านเว็บแอปพลิเคชันที่รองรับการกรองตามแบรนด์และแสดงเหตุผลของการแนะนำ"
)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
set_font(p.add_run("คำสำคัญ: "), size=15, bold=True)
set_font(p.add_run("ระบบแนะนำ, Content-Based Filtering, Cosine Similarity, Jaccard Overlap, "
                   "Feature Engineering, Hybrid Scoring, Sparse Matrix, Streamlit"), size=15)

doc.add_paragraph()
add_heading(doc, "Abstract", size=FONT_SIZE_H1)
add_body(doc,
    "This project presents a content-based perfume recommendation system for a dataset "
    "of approximately 26,000 perfumes sourced from doevent/perfume on HuggingFace. "
    "Sparse feature vectors are constructed from ingredients, scent family, and gender "
    "with per-group weights tuned through a systematic weight sweep. "
    "Results are ranked by a hybrid score combining cosine similarity and Jaccard note overlap. "
    "The optimised system achieves a 16% improvement in Family Precision@10 and an 8% improvement "
    "in Leave-One-Out Hit-Rate@10 over the cosine-only baseline. "
    "No user rating data is required. Queries resolve in under 10 ms and are served through "
    "a web application supporting brand filtering and matched-notes explanations."
)
p = doc.add_paragraph()
set_font(p.add_run("Keywords: "), size=15, bold=True)
set_font(p.add_run("Recommendation System, Content-Based Filtering, Cosine Similarity, "
                   "Jaccard Overlap, Feature Engineering, Hybrid Scoring, Sparse Matrix, Streamlit"), size=15)

page_break(doc)

# ── Chapter 1 ────────────────────────────────────────────────────────────────

add_heading(doc, "1. บทนำ")

add_subheading(doc, "1.1 ที่มา และความสำคัญ")
add_body(doc,
    "ปัจจุบันตลาดน้ำหอมมีผลิตภัณฑ์หลากหลายนับหมื่นรายการ การค้นหาน้ำหอมที่ตรงกับความชอบ"
    "ส่วนตัวเป็นเรื่องยากสำหรับผู้บริโภคทั่วไป เนื่องจากต้องอาศัยความรู้เฉพาะทางด้านส่วนผสม "
    "กลุ่มกลิ่น และการผสมผสานระหว่างโน้ตต่างๆ ระบบแนะนำแบบ Content-Based Filtering "
    "สามารถช่วยค้นพบน้ำหอมที่สอดคล้องกับรสนิยมได้โดยไม่ต้องพึ่งพาประวัติการให้คะแนน "
    "ซึ่งแก้ปัญหา Cold Start Problem ที่พบบ่อยใน Collaborative Filtering ได้ดี (Lops et al., 2011)"
)

add_subheading(doc, "1.2 วัตถุประสงค์")
for i, t in enumerate([
    "พัฒนาระบบแนะนำน้ำหอมแบบ Content-Based ที่ไม่ต้องการข้อมูล User Rating",
    "เปรียบเทียบแนวทางการสร้าง Feature Vector หลายแบบและเลือกแนวทางที่ให้ผลดีที่สุด",
    "ปรับปรุงคุณภาพการแนะนำด้วย Hybrid Scoring และการจูน Weight แบบอัตโนมัติ",
    "พัฒนากระบวนการประเมินผลที่ Reproducible ด้วย Precision@K และ Leave-One-Out Protocol",
    "นำเสนอระบบผ่านเว็บแอปพลิเคชันที่ใช้งานง่ายและอธิบายเหตุผลของการแนะนำได้",
], 1):
    add_numbered(doc, t, i)

add_subheading(doc, "1.3 ทฤษฎีและเทคโนโลยีที่เกี่ยวข้อง")

add_subheading(doc, "1.3.1 Content-Based Filtering", level=3)
add_body(doc,
    "Content-Based Filtering เป็นแนวทางของระบบแนะนำที่วิเคราะห์คุณลักษณะของสินค้า "
    "และสร้าง Item Profile โดยไม่ต้องใช้ข้อมูลพฤติกรรมของผู้ใช้คนอื่น "
    "เหมาะกับ Domain ที่ไม่มี Rating Data และสามารถแก้ปัญหา Cold Start ได้ (Lops et al., 2011)"
)

add_subheading(doc, "1.3.2 Cosine Similarity และ Jaccard Similarity", level=3)
add_body(doc,
    "Cosine Similarity วัดความคล้ายคลึงระหว่าง Vector สองตัว โดยคำนวณค่า Cosine "
    "ของมุมระหว่าง Vector ให้ค่าในช่วง [0, 1] และไม่ขึ้นกับขนาด (Magnitude) ของ Vector "
    "ซึ่งเหมาะกับ Sparse Vector ที่มีส่วนผสมหลากหลายในสัดส่วนต่างกัน "
    "สูตรการคำนวณแสดงดังสมการ (1):"
)
add_equation(doc, "1",
    r"\cos(A,B)=\frac{A\cdot B}{\|A\|\,\|B\|}=\frac{\sum_i A_i B_i}{\sqrt{\sum_i A_i^2}\,\sqrt{\sum_i B_i^2}}"
)
add_body(doc,
    "Jaccard Similarity วัด Overlap ระหว่างสองเซตโดยตรง เหมาะสำหรับการวัด "
    "ว่าส่วนผสมที่ผู้ใช้เลือกปรากฏในน้ำหอมที่แนะนำมากเพียงใด (Jaccard, 1912) "
    "สูตรการคำนวณแสดงดังสมการ (2):"
)
add_equation(doc, "2",
    r"J(A,B)=\frac{|A\cap B|}{|A\cup B|}"
)
add_body(doc,
    "ในโครงงานนี้ใช้ทั้งสอง Metric ร่วมกันในรูปแบบ Hybrid Score "
    "เพื่อให้ผลลัพธ์ที่สมดุลระหว่างความคล้ายคลึงในเชิง Vector Space และการ Match โดยตรงของโน้ต"
)

add_subheading(doc, "1.3.3 Feature Engineering สำหรับข้อมูลน้ำหอม", level=3)
add_body(doc, "การแปลงข้อมูลน้ำหอมให้เป็น Numeric Vector ใช้เทคนิคหลัก 2 ประเภท:")
for t in [
    "Multi-Hot Encoding — แปลง List ส่วนผสมให้เป็น Binary Vector โดยแต่ละมิติแทนส่วนผสมหนึ่งชนิด",
    "One-Hot Encoding — แปลงตัวแปร Categorical เช่น กลุ่มกลิ่น และเพศ ให้เป็น Binary Vector",
    "Rare-Vocabulary Filtering — ตัดส่วนผสมที่พบน้อยกว่าเกณฑ์ขั้นต่ำออกเพื่อลด Noise",
    "Per-Group Feature Weighting — กำหนดน้ำหนักต่างกันสำหรับแต่ละกลุ่มคุณลักษณะแทน Multiplier แบบ Hard-coded",
]:
    add_bullet(doc, t)

add_subheading(doc, "1.3.4 เว็บแอปพลิเคชัน Machine Learning", level=3)
add_body(doc,
    "การนำเสนอระบบผ่านเว็บแอปพลิเคชันช่วยให้ผู้ใช้ทดสอบและโต้ตอบกับโมเดลได้โดยตรง "
    "โดยไม่ต้องมีความรู้ด้านการเขียนโปรแกรม ซึ่งเป็นแนวทางที่นิยมใน Data Science Prototyping"
)

add_subheading(doc, "1.4 ขอบเขตโครงงาน")
for t in [
    "ชุดข้อมูล: doevent/perfume จาก HuggingFace (~26,000 รายการ)",
    "ระบบแนะนำแบบ Content-Based เท่านั้น ไม่รวม Collaborative Filtering",
    "เว็บแอปพลิเคชันสำหรับสาธิตบนเครื่อง Local",
    "ประเมินผลด้วย Intrinsic Metrics เท่านั้น ไม่มี Human Evaluation",
]:
    add_bullet(doc, t)

add_subheading(doc, "1.5 ประโยชน์ที่คาดว่าจะได้รับ")
for i, t in enumerate([
    "ผู้ใช้ค้นหาน้ำหอมที่ตรงรสนิยมได้โดยไม่ต้องมีประวัติการใช้งาน (No Cold Start)",
    "ได้กระบวนการประเมิน Recommendation System แบบ Reproducible สำหรับข้อมูลที่ไม่มี Rating",
    "ได้ต้นแบบสถาปัตยกรรมที่ขยายไปยังโดเมนสินค้าอื่นได้",
], 1):
    add_numbered(doc, t, i)

page_break(doc)

# ── Chapter 2 ────────────────────────────────────────────────────────────────

add_heading(doc, "2. วิธีการดำเนินโครงงาน")

add_subheading(doc, "2.1 การรวบรวมข้อมูล")
add_body(doc,
    "ชุดข้อมูลที่ใช้คือ doevent/perfume จาก HuggingFace Hub "
    "ซึ่งเป็น Open Dataset ที่รวบรวมข้อมูลน้ำหอมกว่า 26,000 รายการ "
    "ดาวน์โหลดอัตโนมัติผ่าน HuggingFace Hub API และจัดเก็บในรูปแบบ JSON "
    "คอลัมน์หลักของชุดข้อมูลแสดงในตารางที่ 1:"
)
make_table(doc,
    ["คุณลักษณะ", "ประเภท", "คำอธิบาย"],
    [
        ["ชื่อแบรนด์ (Brand)", "Categorical", "บริษัทหรือนักออกแบบผู้ผลิต"],
        ["ชื่อน้ำหอม (Name)", "Text", "ชื่อผลิตภัณฑ์"],
        ["กลุ่มกลิ่น (Family)", "Categorical", "หมวดหมู่กลิ่นหลัก เช่น FLORAL, WOODY"],
        ["กลุ่มกลิ่นย่อย (Subfamily)", "Categorical", "หมวดหมู่กลิ่นย่อย"],
        ["ส่วนผสม (Ingredients)", "List of Text", "รายการโน้ตกลิ่น เช่น Rose, Musk, Bergamot"],
        ["เพศเป้าหมาย (Gender)", "Categorical", "MALE, FEMALE หรือ UNISEX"],
        ["คำอธิบาย (Description)", "Free-text", "คำอธิบายลักษณะกลิ่นแบบร้อยแก้ว"],
        ["ภาพผลิตภัณฑ์", "Image Reference", "ชื่ออ้างอิงรูปภาพขวดน้ำหอม"],
    ],
    col_widths=[4.5, 3.0, 8.0]
)
add_caption(doc, "ตารางที่ 1: คุณลักษณะหลักของชุดข้อมูลน้ำหอม")

add_subheading(doc, "2.2 การสำรวจข้อมูลเบื้องต้น")
add_body(doc,
    "ขั้นตอนการสำรวจข้อมูลวิเคราะห์โครงสร้างและการกระจายตัวของชุดข้อมูล "
    "เพื่อทำความเข้าใจก่อนเข้าสู่กระบวนการเตรียมข้อมูล ผลการสำรวจสรุปได้ดังนี้:"
)
for t in [
    "พบน้ำหอมทั้งหมด 26,319 รายการ จาก 1,263 แบรนด์",
    "ส่วนผสมที่พบบ่อยที่สุด 5 อันดับแรก: Musk (12,270 ขวด), Bergamot (9,037), Amber (7,580), Patchouli (7,486), Sandalwood (7,462)",
    "พบส่วนผสมที่ไม่ซ้ำกันทั้งหมด 394 ชนิด ก่อนการกรอง Vocabulary",
    "การกระจายตามเพศ: Female มากที่สุด รองลงมาคือ Unisex และ Male",
    "ไม่พบ Missing Value ในคอลัมน์ Ingredients และ Family",
]:
    add_bullet(doc, t)

add_subheading(doc, "2.3 การเตรียมข้อมูลและสร้าง Feature Vector")
add_body(doc,
    "กระบวนการเตรียมข้อมูลทั้งฝั่ง Training และ Inference ใช้ตรรกะการ Encode เดียวกัน "
    "โดยออกแบบให้โมดูล Feature Engineering เป็นจุดศูนย์กลางของระบบ "
    "เพื่อให้มั่นใจว่า Feature Matrix ที่ใช้ในการสร้างโมเดลและ Query Vector ที่สร้างขณะ Inference "
    "ใช้กระบวนการ Encode แบบเดียวกันเสมอ ซึ่งป้องกันความคลาดเคลื่อนระหว่างสองขั้นตอน"
)

add_subheading(doc, "2.3.1 การกรอง Vocabulary ส่วนผสมที่หายาก (Rare-Vocabulary Filtering)", level=3)
add_body(doc,
    "ส่วนผสมที่ปรากฏในน้ำหอมน้อยกว่า 5 รายการถูกตัดออกจาก Vocabulary "
    "เนื่องจากส่วนผสมที่หายากมากเพิ่ม Noise ใน Feature Space โดยไม่ช่วยในการแนะนำ "
    "ผลการกรอง: ลด Vocabulary จาก 394 เหลือ 312 ส่วนผสม (ตัดออก 82 รายการ) "
    "โดยมีน้ำหอมที่สูญเสียส่วนผสมทั้งหมดเพียง 3 รายการจาก 26,319"
)

add_subheading(doc, "2.3.2 การสร้างและถ่วงน้ำหนัก Feature Vector", level=3)
add_body(doc,
    "Feature Vector ของน้ำหอมแต่ละขวดสร้างจากการเชื่อม (Concatenation) "
    "ของกลุ่มคุณลักษณะ 4 กลุ่ม โดยแต่ละกลุ่มได้รับการถ่วงน้ำหนักตามความสำคัญ "
    "ดังแสดงในตารางที่ 2 ค่า Weight ได้จากการทดสอบ Weight Sweep "
    "โดยเลือกชุดที่ให้ Leave-One-Out Hit-Rate@10 สูงสุด:"
)
make_table(doc,
    ["กลุ่มคุณลักษณะ", "วิธี Encoding", "จำนวนมิติ", "Weight"],
    [
        ["ส่วนผสม (Ingredients)", "Multi-Hot Encoding", "312", "1.00"],
        ["กลุ่มกลิ่น (Family)", "One-Hot Encoding", "~15", "0.50"],
        ["กลุ่มกลิ่นย่อย (Subfamily)", "One-Hot Encoding", "~8", "0.30"],
        ["เพศเป้าหมาย (Gender)", "One-Hot Encoding", "~4", "0.20"],
    ],
    col_widths=[4.5, 4.5, 3.0, 2.5]
)
add_caption(doc, "ตารางที่ 2: องค์ประกอบและน้ำหนักของ Feature Vector (รวม 337 มิติ)")

add_subheading(doc, "2.4 สถาปัตยกรรมระบบ")
add_body(doc,
    "ระบบแบ่งออกเป็น 2 Pipeline หลัก ได้แก่ Offline Pipeline "
    "ที่ดำเนินการก่อนการใช้งาน (สร้าง Encoder และ Feature Matrix) "
    "และ Online Pipeline ที่ทำงานขณะที่ผู้ใช้ส่ง Query "
    "สถาปัตยกรรมโดยรวมแสดงในภาพที่ 1:"
)
doc.add_picture(DIAG_PATH, width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_caption(doc, "ภาพที่ 1: สถาปัตยกรรมระบบ แสดง Offline Pipeline (ซ้าย) และ Online Pipeline (ขวา)")

add_subheading(doc, "2.5 Hybrid Scoring", level=2)
add_body(doc,
    "แทนที่จะใช้ Cosine Similarity เพียงอย่างเดียว ระบบจัดอันดับผลลัพธ์ด้วย Hybrid Score "
    "ที่ผสาน Cosine Similarity กับ Jaccard Overlap ดังสมการ (3) "
    "ซึ่งช่วยให้น้ำหอมที่ Notes ตรงกับที่ผู้ใช้เลือกถูกจัดอันดับสูงขึ้น "
    "นอกจากนี้ยังสามารถแสดง Matched Notes เพื่ออธิบายเหตุผลของการแนะนำได้:"
)
add_equation(doc, "3",
    r"\mathrm{hybrid\_score}=0.70\cdot\cos(Q,P)+0.30\cdot J(Q_{\mathrm{notes}},P_{\mathrm{notes}})",
    width=Inches(6.2),
)
add_body(doc,
    "โดยที่ Q คือ Query Vector ของผู้ใช้, P คือ Feature Vector ของน้ำหอมแต่ละขวด, "
    "Q_notes คือเซตส่วนผสมที่ผู้ใช้เลือก, P_notes คือเซตส่วนผสมของน้ำหอม "
    "ค่าสัมประสิทธิ์ 0.70 และ 0.30 ได้มาจากการทดลองเชิงประจักษ์"
)

add_subheading(doc, "2.6 การพัฒนาเว็บแอปพลิเคชัน")
add_body(doc, "ระบบนำเสนอผ่านเว็บแอปพลิเคชัน Machine Learning รองรับฟีเจอร์ต่อไปนี้:")
for t in [
    "ตัวกรองเพศ กลุ่มกลิ่น กลุ่มกลิ่นย่อย และแบรนด์ผ่าน Dropdown",
    "Multiselect ส่วนผสมจาก 200 ส่วนผสมที่พบบ่อยที่สุด",
    "แสดงผลเป็น Card — ชื่อน้ำหอม แบรนด์ ส่วนผสม Matched Notes และ Hybrid Score เป็นเปอร์เซ็นต์",
    "แสดงรูปภาพผลิตภัณฑ์ (เมื่อมีการดาวน์โหลด Image Dataset)",
    "ระบบ Cache ผลลัพธ์ในหน่วยความจำสำหรับ Query ที่ซ้ำกัน",
]:
    add_bullet(doc, t)

page_break(doc)

# ── Chapter 3 ────────────────────────────────────────────────────────────────

add_heading(doc, "3. ผลการดำเนินงาน")

add_subheading(doc, "3.1 ผลลัพธ์จากกระบวนการเตรียมโมเดล")
add_body(doc,
    "หลังจากผ่านกระบวนการ Preprocessing ระบบสร้างองค์ประกอบหลักสำหรับการแนะนำ "
    "ดังแสดงในตารางที่ 3:"
)
make_table(doc,
    ["องค์ประกอบ", "คำอธิบาย", "ขนาด (โดยประมาณ)"],
    [
        ["ตัว Encode ส่วนผสม (Trained)", "Multi-Hot Encoder สำหรับส่วนผสม 312 ชนิด", "~13 KB"],
        ["ตัว Encode หมวดหมู่ (Trained)", "One-Hot Encoder สำหรับ Family, Subfamily, Gender", "~5 KB"],
        ["ฐานข้อมูลน้ำหอมที่ทำความสะอาดแล้ว", "DataFrame ที่ผ่านการ Normalize (~26,319 แถว)", "~19 MB"],
        ["Feature Matrix แบบ Sparse (Weighted)", "เมทริกซ์คุณลักษณะถ่วงน้ำหนัก ขนาด 26,319 × 337", "~340 KB"],
        ["ตัวระบุ Configuration ของโมเดล", "ค่า Weight, Hybrid Blend และขนาด Vocabulary", "<1 KB"],
    ],
    col_widths=[5.5, 7.0, 4.0]
)
add_caption(doc, "ตารางที่ 3: องค์ประกอบของโมเดลที่ได้จากกระบวนการ Build")

add_subheading(doc, "3.2 การเปรียบเทียบแนวทาง Feature Vector (Legacy Comparison)")
add_body(doc,
    "ระบบเบื้องต้นทดสอบแนวทางการสร้าง Feature Vector 4 แบบ (Approach A–D) "
    "บน 5 Test Queries ที่เตรียมไว้ล่วงหน้า โดยใช้ 3 Intrinsic Metrics ดังนี้:"
)
for t in [
    "Ingredient Overlap — ค่าเฉลี่ย Jaccard Similarity ระหว่างส่วนผสมของ Query และผลลัพธ์ที่แนะนำ",
    "Family Match Rate — สัดส่วนของผลลัพธ์ที่มีกลุ่มกลิ่นตรงกับ Query",
    "Intra-List Diversity (ILD) — ค่าเฉลี่ย Pairwise Cosine Distance ภายในผลลัพธ์ (สูง = หลากหลาย)",
]:
    add_bullet(doc, t)
add_body(doc,
    "คะแนนรวม (Avg Score) คำนวณเป็นค่าเฉลี่ยเลขคณิตของ Metric ทั้ง 3 "
    "ดังสมการ (4):"
)
add_equation(doc, "4",
    r"\mathrm{Avg\ Score}=\frac{\mathrm{Ingredient\ Overlap}+\mathrm{Family\ Match\ Rate}+\mathrm{ILD}}{3}",
    width=Inches(6.5),
)
make_table(doc,
    ["แนวทาง", "Feature Vector", "Ingredient Overlap↑", "Family Match↑", "ILD↑", "Avg Score↑"],
    [
        ["A", "Multi-Hot Encoding (ส่วนผสมอย่างเดียว)", "0.5912", "0.360", "0.8612", "0.6041"],
        ["B", "ส่วนผสม + ข้อมูลหมวดหมู่", "0.4999", "0.580", "0.8197", "0.6332"],
        ["C", "B + TF-IDF จากคำอธิบายน้ำหอม (Salton & Buckley, 1988)", "0.4937", "0.580", "0.8308", "0.6348"],
        ["D", "ส่วนผสมน้ำหนัก 2× + ข้อมูลหมวดหมู่", "0.5793", "0.480", "0.8479", "0.6357"],
    ],
    col_widths=[1.8, 5.2, 2.8, 2.5, 1.8, 2.5]
)
add_caption(doc, "ตารางที่ 4: ผลการเปรียบเทียบ Feature Vector 4 แนวทาง (↑ = ยิ่งสูงยิ่งดี)")
add_body(doc,
    "แนวทาง D ให้ Avg Score สูงสุด (0.6357) เนื่องจากการเน้นส่วนผสมเป็นสัญญาณหลัก "
    "จึงถูกเลือกเป็น Baseline ก่อนเข้าสู่ขั้นตอนการ Optimize"
)

add_subheading(doc, "3.3 ผลการ Optimize ด้วย Hybrid Model (ระบบปัจจุบัน)")

add_subheading(doc, "3.3.1 ผลการทดสอบ Weight Sweep", level=3)
add_body(doc,
    "กระบวนการประเมินผลทดสอบชุด Weight หลายชุดโดยวัดด้วย Leave-One-Out Hit-Rate@10 "
    "บน 200 Query ตัวอย่าง ผลแสดงในตารางที่ 5:"
)
make_table(doc,
    ["น้ำหนักส่วนผสม", "น้ำหนัก Family", "น้ำหนัก Subfamily", "น้ำหนัก Gender", "LOO Hit-Rate@10"],
    [
        ["1.0", "0.2", "0.1", "0.1", "0.7602"],
        ["1.0", "0.3", "0.2", "0.15", "0.7839"],
        ["1.0", "0.5", "0.3", "0.2", "0.8520 (ดีที่สุด)"],
        ["2.0", "0.3", "0.2", "0.15", "0.6979"],
    ],
    col_widths=[3.0, 3.0, 3.0, 3.0, 3.5]
)
add_caption(doc, "ตารางที่ 5: ผลการทดสอบ Weight Sweep")
add_body(doc,
    "ชุด Weight ที่เน้น Family/Subfamily/Gender มากขึ้น (0.5/0.3/0.2) "
    "ให้ Leave-One-Out Hit-Rate@10 สูงสุดที่ 0.852 "
    "การเพิ่มน้ำหนัก Family แทนการดัน Ingredients แบบ Hard-coded (Approach D เดิม) "
    "ทำให้โมเดลเรียนรู้ว่า Category Context สำคัญต่อการระบุน้ำหอมที่ถูกต้อง"
)

add_subheading(doc, "3.3.2 การเปรียบเทียบก่อนและหลัง Optimize", level=3)
add_body(doc,
    "ตารางที่ 6 เปรียบเทียบผลการประเมินก่อนและหลัง Optimize "
    "วัดบน Protocol เดียวกัน (300 Query ตัวอย่าง, k=10) "
    "เพื่อให้เปรียบเทียบได้ตรงๆ:"
)
make_table(doc,
    ["Metric", "ความหมาย", "ก่อน Optimize", "หลัง Optimize", "เปลี่ยน (Δ)"],
    [
        ["Precision@10 (Family)",
         "สัดส่วนผลลัพธ์ที่กลุ่มกลิ่นตรงกับที่ผู้ใช้เลือก",
         "0.545", "0.634", "+0.089 (+16%)"],
        ["Precision@10 (≥2 Notes)",
         "สัดส่วนผลลัพธ์ที่มีโน้ตร่วมกับ Query ≥2 ชนิด",
         "0.943", "0.937", "−0.006 (−0.6%)"],
        ["Leave-One-Out Hit-Rate@10",
         "อัตราการดึงน้ำหอมที่ซ่อนไว้กลับใน Top-10",
         "0.780", "0.844", "+0.064 (+8%)"],
    ],
    col_widths=[4.0, 5.0, 3.0, 3.0, 3.5]
)
add_caption(doc, "ตารางที่ 6: เปรียบเทียบ Metric ก่อนและหลัง Optimize")
add_body(doc,
    "Family Precision และ Hit-Rate ดีขึ้นชัดเจน ส่วน Note Precision ลดนิดหน่อย (−0.6%) "
    "เนื่องจาก Hybrid Scoring ให้ Category มีบทบาทมากขึ้นใน Ranking "
    "ถือเป็น Trade-off ที่ยอมรับได้ เพราะยังคง Match โน้ตได้ ~94% ของผลลัพธ์"
)

add_subheading(doc, "3.4 ประสิทธิภาพของระบบ (System Performance)")
make_table(doc,
    ["ด้าน", "รายละเอียด"],
    [
        ["ขนาด Catalog", "26,319 น้ำหอม"],
        ["ขนาด Feature Matrix", "26,319 × 337 มิติ (Sparsity ~96.82%)"],
        ["เวลาตอบสนองต่อ Query", "< 10 ms ต่อ 1 Query (ไม่มี Cache)"],
        ["เวลาตอบสนองเมื่อ Query ซ้ำ", "< 1 ms (ระบบ Cache ในหน่วยความจำ สูงสุด 256 รายการ)"],
        ["การใช้ Memory", "~150–200 MB (Feature Matrix + ฐานข้อมูลใน RAM)"],
        ["Hardware ที่ต้องการ", "CPU ล้วน ไม่ต้องใช้ GPU"],
        ["เทคโนโลยีหลัก", "Python 3.11, scikit-learn ≥1.3, scipy, Streamlit"],
    ],
    col_widths=[5.5, 11.0]
)
add_caption(doc, "ตารางที่ 7: ประสิทธิภาพของระบบ")

add_subheading(doc, "3.5 ปัญหาที่พบและแนวทางแก้ไข (Issues and Solutions)")
make_table(doc,
    ["ปัญหา", "ผลกระทบ", "แนวทางแก้ไข"],
    [
        ["ส่วนผสมที่ไม่อยู่ใน Vocabulary",
         "Query Vector เป็นศูนย์สำหรับโน้ตนั้น → Match อ่อน",
         "กรอง Vocabulary ให้ครอบคลุม ส่วนผสมหลักและบันทึก Warning"],
        ["ส่วนผสมที่พบน้อยมาก (Long-tail)",
         "เพิ่ม Noise ใน Feature Space",
         "กำหนดเกณฑ์ความถี่ขั้นต่ำ ตัดส่วนผสมหายาก 82 ชนิดออก"],
        ["การสร้างโมเดลใหม่ซ้ำกันทุก Request",
         "เวลา Latency สูง",
         "เปลี่ยนเป็นการคำนวณ Cosine Similarity แบบ Sparse รอบเดียว"],
        ["ความคลาดเคลื่อนระหว่าง Training/Inference Encoding",
         "Feature Matrix กับ Query Vector ต่างกัน",
         "ออกแบบให้ทั้งสองขั้นตอนใช้โมดูล Feature Engineering เดียวกัน"],
        ["ข้อมูลภาพผลิตภัณฑ์ขนาดใหญ่ (~835 MB)",
         "ไม่เหมาะกับการจัดเก็บใน Version Control",
         "ออกแบบให้ดาวน์โหลดแยกผ่าน UI ของแอปพลิเคชัน"],
    ],
    col_widths=[4.5, 5.0, 6.5]
)
add_caption(doc, "ตารางที่ 8: ปัญหาที่พบและแนวทางแก้ไข")

page_break(doc)

# ── Chapter 4 ────────────────────────────────────────────────────────────────

add_heading(doc, "4. สรุปและข้อเสนอแนะ")

add_subheading(doc, "4.1 สรุปผล")
add_body(doc,
    "โครงงานนี้พัฒนาระบบแนะนำน้ำหอมแบบ Content-Based Filtering ที่ไม่ต้องการ User Rating "
    "โดยใช้ Hybrid Scoring (Cosine + Jaccard) บน Sparse Feature Vector "
    "ที่ได้รับการ Optimize ด้วยการกรอง Vocabulary ที่หายาก การถ่วงน้ำหนักต่อกลุ่มคุณลักษณะ "
    "และการจูน Weight แบบอัตโนมัติ ผลการประเมินยืนยันว่าการปรับปรุงทำให้ "
    "Family Precision@10 เพิ่มขึ้น 16% และ Leave-One-Out Hit-Rate@10 เพิ่มขึ้น 8% "
    "เมื่อเทียบกับ Baseline ในขณะที่ Note Precision ยังคงสูงที่ ~94% "
    "ระบบทำงานบน CPU ล้วน ตอบสนองได้น้อยกว่า 10 ms ต่อ Query "
    "และนำเสนอผ่านเว็บแอปพลิเคชันที่อธิบายผลลัพธ์ด้วยการแสดงโน้ตที่ตรงกัน"
)

add_subheading(doc, "4.2 ข้อจำกัด")
for t in [
    "ไม่มี User Rating ทำให้ไม่สามารถใช้ Collaborative Filtering หรือ Personalization ได้",
    "การประเมินผลอาศัย Intrinsic Metrics เท่านั้น ยังไม่มีการทดสอบกับผู้ใช้จริง",
    "ชุดข้อมูลไม่ได้รับการ Update Real-time ต้องสร้างโมเดลใหม่เมื่อมีข้อมูลเพิ่ม",
    "ข้อมูลคำอธิบาย Free-text ยังไม่ถูกนำมาใช้ใน Production Feature Vector",
]:
    add_bullet(doc, t)

add_subheading(doc, "4.3 ข้อเสนอแนะสำหรับการพัฒนาในอนาคต")
for i, t in enumerate([
    "นำ Sentence-Transformers มาใช้ Encode คำอธิบายน้ำหอม เพื่อเพิ่ม Semantic Similarity",
    "เก็บ Implicit Feedback จากพฤติกรรมผู้ใช้เพื่อพัฒนา Hybrid Recommender",
    "เพิ่ม Maximal Marginal Relevance (MMR) เพื่อเพิ่มความหลากหลายของผลลัพธ์",
    "พัฒนา REST API เพื่อรองรับการใช้งานจากหลาย Platform",
    "ทดสอบ Approximate Nearest Neighbor Search หากขนาด Catalog เติบโตขึ้นมาก",
    "ดำเนินการ Human Evaluation อย่างเป็นระบบเพื่อวัด User Satisfaction จริง",
], 1):
    add_numbered(doc, t, i)

page_break(doc)

# ── References ────────────────────────────────────────────────────────────────

add_heading(doc, "5. เอกสารอ้างอิง")
refs = [
    'HuggingFace, "doevent/perfume dataset," [Online]. Available: https://huggingface.co/datasets/doevent/perfume',
    'scikit-learn developers, "Multi-Hot Encoding: MultiLabelBinarizer," [Online]. Available: https://scikit-learn.org/stable/modules/generated/sklearn.preprocessing.MultiLabelBinarizer.html',
    'scikit-learn developers, "One-Hot Encoding: OneHotEncoder," [Online]. Available: https://scikit-learn.org/stable/modules/generated/sklearn.preprocessing.OneHotEncoder.html',
    'scikit-learn developers, "Cosine Similarity: sklearn.metrics.pairwise.cosine_similarity," [Online]. Available: https://scikit-learn.org/stable/modules/generated/sklearn.metrics.pairwise.cosine_similarity.html',
    'Streamlit Inc., "Streamlit — A faster way to build and share data apps," [Online]. Available: https://streamlit.io/',
    'SciPy Community, "Sparse matrix operations," [Online]. Available: https://docs.scipy.org/doc/scipy/reference/sparse.html',
    'Lops, P., de Gemmis, M., & Semeraro, G. (2011). Content-based Recommender Systems: State of the Art and Trends. In Recommender Systems Handbook (pp. 73–105). Springer.',
    'Salton, G., & Buckley, C. (1988). Term-weighting approaches in automatic text retrieval. Information Processing & Management, 24(5), 513–523.',
    'Jaccard, P. (1912). The distribution of the flora in the alpine zone. New Phytologist, 11(2), 37–50.',
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(f"[{i}] {r}"), size=14)

doc.add_paragraph()

# ── Appendix (technical detail OK here) ─────────────────────────────────────

add_heading(doc, "ภาคผนวก ก — โครงสร้างโปรเจกต์และคำสั่งสำคัญ", size=FONT_SIZE_H2)
add_body(doc,
    "ภาคผนวกนี้แสดงรายละเอียดทางเทคนิคสำหรับนักพัฒนาที่ต้องการทำความเข้าใจ "
    "โครงสร้างไฟล์และวิธีรันระบบ"
)

add_subheading(doc, "ก.1 โครงสร้างโปรเจกต์ (Project Structure)")
for line in [
    "perfume_recommender/",
    "  data/perfumes.json            ← Dataset ดิบ (JSON)",
    "  models/                       ← Artifacts ที่ Build แล้ว",
    "    mlb_ingredients.pkl         ← Trained Multi-Hot Encoder",
    "    ohe_categories.pkl          ← Trained One-Hot Encoder",
    "    perfume_df.pkl              ← Cleaned DataFrame",
    "    matrix_HYBRID.npz           ← Weighted Sparse Feature Matrix",
    "    feature_config.pkl          ← Weight + Vocab Configuration",
    "  notebooks/                    ← EDA + Legacy KNN Comparison",
    "  src/",
    "    feature_engineering.py      ← Shared Encoding Logic",
    "    build_models.py             ← Build Pipeline Script",
    "    evaluate.py                 ← Evaluation Script",
    "    recommender.py              ← Recommendation Engine",
    "    app.py                      ← Streamlit Web Application",
]:
    add_code(doc, line)

add_subheading(doc, "ก.2 คำสั่งสำคัญ (Key Commands)")
for label, cmd in [
    ("ติดตั้ง dependencies",   "pip install -r requirements.txt"),
    ("ดาวน์โหลด dataset",      "python -c \"from src.data_loader import download_dataset; download_dataset()\""),
    ("สร้าง Model Artifacts",  "python src/build_models.py"),
    ("ประเมินโมเดล",           "python src/evaluate.py"),
    ("รัน Web Application",    "streamlit run src/app.py"),
]:
    p = doc.add_paragraph()
    set_font(p.add_run(f"{label}: "), size=14, bold=True)
    set_font(p.add_run(cmd), name="Courier New", size=13)
    p.paragraph_format.space_after = Pt(3)

doc.save(OUT)

# Clean up temp files
if os.path.exists(DIAG_PATH):
    os.remove(DIAG_PATH)
if os.path.isdir(EQ_DIR):
    shutil.rmtree(EQ_DIR, ignore_errors=True)

print(f"Saved: {OUT}")
